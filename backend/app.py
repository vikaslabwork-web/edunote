from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from transformers import pipeline
from docx import Document
from PyPDF2 import PdfReader
from youtube_transcript_api import YouTubeTranscriptApi
import os
import tempfile

app = Flask(__name__)
CORS(app)

# Use the port Render provides
PORT = int(os.environ.get("PORT", 5000))

# Initialize summarization pipeline (v4.34.0)
summarizer = pipeline("summarization", model="sshleifer/distilbart-cnn-12-6")

# Max file upload size 16 MB
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

# ------------------- Helper Functions -------------------

def summarize_text(text):
    summary = summarizer(text, max_length=150, min_length=30, do_sample=False)
    return summary[0]['summary_text']

def extract_text_from_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def extract_text_from_youtube(video_url):
    video_id = video_url.split("v=")[-1]
    transcript_list = YouTubeTranscriptApi.get_transcript(video_id)
    text = " ".join([t['text'] for t in transcript_list])
    return text

# ------------------- Routes -------------------

@app.route("/summarize-text", methods=["POST"])
def summarize_text_route():
    data = request.json
    text = data.get("text", "")
    if not text:
        return jsonify({"error": "No text provided"}), 400
    summary = summarize_text(text)
    return jsonify({"summary": summary})

@app.route("/summarize-file", methods=["POST"])
def summarize_file_route():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    
    file = request.files["file"]
    filename = file.filename.lower()
    temp_file = tempfile.NamedTemporaryFile(delete=False)
    file.save(temp_file.name)

    if filename.endswith(".pdf"):
        text = extract_text_from_pdf(temp_file.name)
    elif filename.endswith(".docx"):
        text = extract_text_from_docx(temp_file.name)
    else:
        return jsonify({"error": "Unsupported file type"}), 400

    summary = summarize_text(text)
    return jsonify({"summary": summary})

@app.route("/summarize-youtube", methods=["POST"])
def summarize_youtube_route():
    data = request.json
    url = data.get("url", "")
    if not url:
        return jsonify({"error": "No YouTube URL provided"}), 400
    
    try:
        text = extract_text_from_youtube(url)
        summary = summarize_text(text)
        return jsonify({"summary": summary})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# ------------------- File Download Routes -------------------

@app.route("/download-docx", methods=["POST"])
def download_docx():
    data = request.json
    text = data.get("text", "")
    if not text:
        return jsonify({"error": "No text provided"}), 400
    
    doc = Document()
    doc.add_paragraph(text)
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return send_file(temp_file.name, as_attachment=True, download_name="summary.docx")

@app.route("/download-pdf", methods=["POST"])
def download_pdf():
    from fpdf import FPDF
    data = request.json
    text = data.get("text", "")
    if not text:
        return jsonify({"error": "No text provided"}), 400
    
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    for line in text.split("\n"):
        pdf.multi_cell(0, 10, line)
    
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    pdf.output(temp_file.name)
    return send_file(temp_file.name, as_attachment=True, download_name="summary.pdf")

# ------------------- Main -------------------

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=PORT)
